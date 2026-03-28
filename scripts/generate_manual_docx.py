import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """
    Set cell border
    Usage: set_cell_border(cell, top={"sz": 12, "color": "#FF0000", "val": "single"}, ...)
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcPr.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcPr.append(element)
            for key, val in edge_data.items():
                element.set(qn('w:{}'.format(key)), str(val))

def create_manual():
    doc = Document()

    # --- PORTADA ---
    title = doc.add_heading('SISTEMA PACT (PROGRAMA ACT)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('MANUAL DE USUARIO OFICIAL Y DETALLADO')
    subtitle_fmt = subtitle.paragraph_format
    subtitle_fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(18)
    subtitle.runs[0].font.color.rgb = RGBColor(37, 99, 235) # Blue

    doc.add_paragraph('\n' * 5)
    
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('Versión: 3.1.5 (Actualizada con IA y Explorador de Archivos)\nFecha: Marzo 2026\nAutoridad de Carreteras y Transportación (ACT)')
    run.font.italic = True
    
    doc.add_page_break()

    # --- TABLA DE CONTENIDO (SIMULADA) ---
    doc.add_heading('TABLA DE CONTENIDO', 1)
    sections = [
        "1. Registro e Inicio de Sesión",
        "2. Dashboard General",
        "3. Estructura de Proyecto (Secciones 1 a 10)",
        "4. Asistente IA de Voz (Smart Dictation)",
        "5. Gestor de Archivos (Proyectos)",
        "6. Generación de Reportes (ACT-45, 96, CCML)",
        "7. Panel de Administración y Accesos",
        "8. Preguntas Frecuentes y Solución de Problemas"
    ]
    for s in sections:
        p = doc.add_paragraph(s)
        p.style = 'List Bullet'

    doc.add_page_break()

    # --- SECCIÓN 1: REGISTRO ---
    doc.add_heading('1. Registro e Inicio de Sesión', 1)
    doc.add_paragraph("Para utilizar PACT, el usuario debe seguir estos pasos:")
    
    p = doc.add_paragraph()
    p.add_run("Paso A: Registro de Solicitud de Acceso").bold = True
    doc.add_paragraph("Al abrir la aplicación por primera vez o si no tiene cuenta, haga clic en 'Solicitar Acceso'. Deberá completar su nombre, email, rol deseado (Ver tabla de niveles) y los números de proyecto (ej: AC-123456) a los que necesita entrar.")
    
    p = doc.add_paragraph()
    p.add_run("Paso B: Aprobación del Administrador").bold = True
    doc.add_paragraph("Un administrador global revisará su solicitud. Una vez aprobada, recibirá un correo electrónico automático con su contraseña temporal.")
    
    p = doc.add_paragraph()
    p.add_run("Paso C: Login").bold = True
    doc.add_paragraph("Use su email y la clave recibida para entrar. El sistema recordará su sesión en esa computadora.")

    # TABLA DE ROLES
    doc.add_heading('Niveles de Acceso (Roles)', 2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Nivel'
    hdr_cells[1].text = 'Nombre'
    hdr_cells[2].text = 'Capacidades'
    
    roles = [
        ('A', 'Súper Admin', 'Acceso total a todos los proyectos, gestión de usuarios y borrado de registros.'),
        ('B', 'Admin Proy.', 'Puede editar datos del proyecto, agregar partidas y gestionar miembros del mismo.'),
        ('C', 'Data Entry', 'Puede agregar y editar partidas, pagos y documentos de cumplimiento.'),
        ('D', 'Lectura', 'Solo puede ver la información y descargar reportes. No puede editar.'),
        ('E', 'Inspector', 'Enfocado en Informes Diarios e Inspecciones. Restringido en áreas contables.')
    ]
    for r, name, desc in roles:
        row_cells = table.add_row().cells
        row_cells[0].text = r
        row_cells[1].text = name
        row_cells[2].text = desc

    # --- SECCIÓN 4: ASISTENTE IA ---
    doc.add_page_break()
    doc.add_heading('4. Asistente IA de Voz (Smart Dictation)', 1)
    doc.add_paragraph("PACT incluye ahora un motor de Inteligencia Artificial para agilizar el trabajo de campo de los inspectores.")
    
    p = doc.add_paragraph()
    p.add_run("¿Cómo funciona?").bold = True
    doc.add_paragraph("En las secciones de 'Informe de Actividades' e 'Inspección', encontrará un botón azul con un icono de micrófono llamado 'Smart Dictation'.")
    
    items_ia = [
        "Presione el botón para iniciar la grabación.",
        "Hable de forma natural: Mencione el personal en obra, el equipo utilizado, las partidas trabajadas con sus cantidades, cualquier demora y anotaciones de seguridad.",
        "Presione 'Detener' (Botón Rojo).",
        "La IA procesará su voz y automáticamente llenará las tablas correspondientes de personal, equipo y notas, clasificando la información por usted."
    ]
    for item in items_ia:
        p = doc.add_paragraph(item, style='List Number')

    # --- SECCIÓN 5: ARCHIVOS ---
    doc.add_heading('5. Gestor de Archivos (Proyectos)', 1)
    doc.add_paragraph("Cada proyecto cuenta con una pestaña de 'Archivos' que emula el Explorador de Windows.")
    
    doc.add_paragraph("Características:", style='List Bullet')
    doc.add_paragraph("Creación de Folders ilimitados para organizar planos, fotos y facturas.", style='List Bullet')
    doc.add_paragraph("Subida de archivos (Drag & Drop) permitiendo cargar múltiples documentos a la vez.", style='List Bullet')
    doc.add_paragraph("Vista Previa: Haga clic en un archivo para verlo directamente en el programa si es PDF o Imagen.", style='List Bullet')
    doc.add_paragraph("Seguridad: Los archivos se almacenan en la nube vinculados únicamente a ese proyecto.", style='List Bullet')

    # --- SECCIÓN 6: REPORTES ---
    doc.add_page_break()
    doc.add_heading('6. Generación de Reportes', 1)
    doc.add_paragraph("El sistema automatiza la creación de formularios oficiales:")
    
    p = doc.add_paragraph()
    p.add_run("ACT-45 (Certificación de Pago):").bold = True
    doc.add_paragraph("Genera el PDF oficial para contabilidad basado en las partidas ejecutadas en la sección 6.")
    
    p = doc.add_paragraph()
    p.add_run("ACT-96 (Registro de Órdenes de Cambio):").bold = True
    doc.add_paragraph("Documento legal que resume los incrementos o decrementos autorizados.")
    
    p = doc.add_paragraph()
    p.add_run("CCML (Contract Change Modification Log):").bold = True
    doc.add_paragraph("Reporte especializado en Excel que rastrea toda la vida financiera del contrato.")

    # --- SECCIÓN 10: BOTONES ---
    doc.add_heading('Manual de Botones de Interacción', 1)
    
    btns = [
        ('Guardar / Sincronizar', 'Ubicado arriba a la derecha. Sube tus cambios locales a la base de datos central.'),
        ('Fijar (Pin)', 'Permite anclar una fila de la tabla para que no se mueva mientras navegas por otras secciones.'),
        ('Zafacón Rojo', 'Borra definitivamente el registro (Partida, Persona o Pago). Requiere confirmación.'),
        ('Smart Dictation', 'Activa el procesamiento de voz por IA.'),
        ('Botones Flotantes (+)', 'Acceso rápido para añadir registros sin tener que subir al inicio de la página.')
    ]
    
    table_btn = doc.add_table(rows=1, cols=2)
    table_btn.style = 'Table Grid'
    h = table_btn.rows[0].cells
    h[0].text = 'Botón / Icono'
    h[1].text = 'Función'
    
    for b, f in btns:
        r = table_btn.add_row().cells
        r[0].text = b
        r[1].text = f

    # --- FINAL ---
    doc.add_paragraph('\n' * 3)
    p_end = doc.add_paragraph("Para soporte técnico contacte al Administrador del Sistema PACT.")
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save('Manual_de_Usuario_ACT_Oficial_V3.docx')
    print("Manual generado exitosamente: Manual_de_Usuario_ACT_Oficial_V3.docx")

if __name__ == "__main__":
    create_manual()
