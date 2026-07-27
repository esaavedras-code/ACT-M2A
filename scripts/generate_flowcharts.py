import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_callout(doc, title, text, bg_hex="F0F4F8", border_hex="1F4E78"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>\n'
        f'  <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_hex}"/>\n'
        f'  <w:top w:val="none"/>\n'
        f'  <w:right w:val="none"/>\n'
        f'  <w:bottom w:val="none"/>\n'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run_t = p.add_run(f"📌 {title}\n")
    run_t.bold = True
    run_t.font.name = "Arial"
    run_t.font.size = Pt(11)
    run_t.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    
    run_b = p.add_run(text)
    run_b.font.name = "Arial"
    run_b.font.size = Pt(10)
    run_b.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(0)
    p_space.paragraph_format.space_after = Pt(6)

def create_flowcharts():
    output_dir = r"C:\Users\Enrique Saavedra\Documents\PROGRAMAS AI\Programa ACT\Documentos\Flowcharts"
    os.makedirs(output_dir, exist_ok=True)

    # 1. HTML CON TODOS LOS DIAGRAMAS DE FLUJO INTERACTIVOS (Mermaid.js + Estilo Premium)
    html_content = """<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Diagramas de Flujo de Procesos - Sistema PACT</title>
    <script src="https://cdn.jsdelivr.net/npm/mermaid@10/dist/mermaid.min.js"></script>
    <style>
        :root {
            --primary: #1F4E78;
            --secondary: #2F5597;
            --bg-dark: #0F172A;
            --card-bg: #1E293B;
            --text-light: #F8FAFC;
            --accent: #38BDF8;
        }
        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            background-color: #0B0F19;
            color: #E2E8F0;
            margin: 0;
            padding: 40px 20px;
        }
        .container {
            max-width: 1200px;
            margin: 0 auto;
        }
        .header-box {
            background: linear-gradient(135deg, #1E293B 0%, #0F172A 100%);
            border: 1px solid #334155;
            border-left: 6px solid #38BDF8;
            border-radius: 12px;
            padding: 30px;
            margin-bottom: 40px;
            box-shadow: 0 10px 25px rgba(0,0,0,0.5);
        }
        h1 {
            color: #F8FAFC;
            font-size: 2.2rem;
            margin: 0 0 10px 0;
        }
        .subtitle {
            color: #94A3B8;
            font-size: 1.1rem;
            margin-bottom: 20px;
        }
        .author-tag {
            display: inline-block;
            background: #1E3A8A;
            color: #93C5FD;
            padding: 6px 14px;
            border-radius: 20px;
            font-size: 0.9rem;
            font-weight: 600;
        }
        .section-card {
            background-color: #1E293B;
            border: 1px solid #334155;
            border-radius: 12px;
            padding: 30px;
            margin-bottom: 40px;
            box-shadow: 0 4px 15px rgba(0,0,0,0.3);
        }
        .section-card h2 {
            color: #38BDF8;
            font-size: 1.5rem;
            border-bottom: 2px solid #334155;
            padding-bottom: 12px;
            margin-top: 0;
        }
        .description {
            color: #CBD5E1;
            font-size: 0.95rem;
            line-height: 1.6;
            margin-bottom: 25px;
        }
        .mermaid {
            background-color: #0F172A;
            padding: 25px;
            border-radius: 8px;
            border: 1px solid #334155;
            overflow-x: auto;
        }
    </style>
</head>
<body>
    <div class="container">
        <div class="header-box">
            <h1>SISTEMA PACT - DIAGRAMAS DE FLUJO DE PROCESOS (FLOWCHARTS)</h1>
            <div class="subtitle">Flujos de Trabajo Operativos, Control Presupuestario y Auditoría Gubernamental</div>
            <div class="author-tag">Diseñado por: Ing. Enrique Saavedra Sada, PE</div>
        </div>

        <!-- 1. RESUMEN -->
        <div class="section-card">
            <h2>1. SECCIÓN: RESUMEN (DASHBOARD & CONSOLIDACIÓN EJECUTIVA)</h2>
            <div class="description">
                El módulo de Resumen centraliza los datos maestros del contrato, consolidando el progreso físico de campo (ACT-45), 
                los desembolsos acumulados (ACT-117C) y las modificaciones contractuales (CHO) para alimentar los indicadores ejecutivos (KPIs) 
                y el informe maestro ACT-32.
            </div>
            <div class="mermaid">
            graph TD
                A[Inicio: Carga de Datos del Proyecto] --> B[Lectura de Contrato Original & CHOs Aprobadas]
                B --> C[Consolidación de Registros Acumulados]
                C --> D1[Progreso Físico en Campo - ACT-45]
                C --> D2[Histórico de Certificaciones - ACT-117C]
                C --> D3[Balance de Retenciones & MOS - ACT-117B]
                D1 --> E[Cálculo de Indicadores KPIs en Tiempo Real]
                D2 --> E
                D3 --> E
                E --> F[Generación de Tableros Visuales & Dashboard]
                F --> G[Exportación de Reporte Ejecutivo Maestro ACT-32]
                G --> H[Fin: Monitoreo & Auditoría Continua]
                
                style A fill:#1F4E78,stroke:#38BDF8,color:#fff
                style E fill:#065F46,stroke:#10B981,color:#fff
                style G fill:#1E3A8A,stroke:#60A5FA,color:#fff
            </div>
        </div>

        <!-- 2. ENTRADA DE DATOS -->
        <div class="section-card">
            <h2>2. SECCIÓN: ENTRADA DE DATOS (INFORMES DIARIOS ACT-45 E INSPECCIÓN ACT-96)</h2>
            <div class="description">
                Gobierna la recolección de evidencia primaria en el sitio de construcción. Captura las jornadas laborales, condiciones 
                climáticas, recursos consumidos, cantidades ejecutadas por partida y hallazgos ambientales/seguridad (ACT-96).
            </div>
            <div class="mermaid">
            graph TD
                A[Inicio Jornada de Inspección] --> B[Crear Nuevo Informe Diario ACT-45]
                B --> C[Registrar Datos Ambientales, Mano de Obra & Equipos]
                C --> D[Seleccionar Partidas Contractuales Activas]
                D --> E[Ingresar Cantidades Físicas Instaladas en el Día]
                E --> F{¿Ocurrió Incidencia / Inspección Específica?}
                F -- Sí --> G[Abrir & Llenar Formulario ACT-96]
                G --> H[Vincular Hallazgo u Orden de Corrección]
                F -- No --> I[Adjuntar Evidencia Fotográfica Anclada]
                H --> I
                I --> J[Firma Digital del Inspector de Campo]
                J --> K[Actualización Automática del Progreso en Campo en PACT]
                K --> L[Fin: Registro Validado en BD]

                style A fill:#1F4E78,stroke:#38BDF8,color:#fff
                style F fill:#D97706,stroke:#FBBF24,color:#fff
                style K fill:#065F46,stroke:#10B981,color:#fff
            </div>
        </div>

        <!-- 3. CUMPLIMIENTO LABORAL -->
        <div class="section-card">
            <h2>3. SECCIÓN: CUMPLIMIENTO LABORAL (CERTIFIED PAYROLL & DAVIS-BACON)</h2>
            <div class="description">
                Módulo fiscalizador de la legislación laboral federal (Davis-Bacon Act) y estatal. Valida las clasificaciones de puestos, 
                tarifas salariales por hora, pago de horas extra y retenciones de nómina de contratistas y subcontratistas.
            </div>
            <div class="mermaid">
            graph TD
                A[Recepción de Nómina Semanal Contratista] --> B[Carga de Certificación de Nómina en PACT]
                B --> C[Auditoría de Clasificaciones de Puestos de Trabajo]
                C --> D[Verificación de Tarifa Horaria vs. Scale de Salarios Mínimos]
                D --> E{¿Cumple con la Tarifa Salarial Obligatoria?}
                E -- No --> F[Generar Alerta de Incumplimiento & Requerimiento de Restitución]
                F --> G[Bloquear Aprobación de Cumplimiento Laboral]
                E -- Sí --> H[Auditar Horas Regulares & Horas Extra (Overtime)]
                H --> I[Validar Declaración Jurada del Patrono]
                I --> J[Aprobar Certificación Laboral Semanal en PACT]
                J --> K[Fin: Habilitar Avance para Certificación de Pago]

                style A fill:#1F4E78,stroke:#38BDF8,color:#fff
                style E fill:#D97706,stroke:#FBBF24,color:#fff
                style F fill:#991B1B,stroke:#F87171,color:#fff
                style J fill:#065F46,stroke:#10B981,color:#fff
            </div>
        </div>

        <!-- 4. CERTIFICADOS DE MANUFACTURA -->
        <div class="section-card">
            <h2>4. SECCIÓN: CERTIFICADOS DE MANUFACTURA (CM & CUMPLIMIENTO BABA)</h2>
            <div class="description">
                Aplica la 'Regla de Oro' de PACT. Asegura que ninguna partida de hierro, acero o producto manufacturado BABA sea certifiable 
                para cobro sin contar con el respaldo documental de procedencia y calidad cargado en el sistema.
            </div>
            <div class="mermaid">
            graph TD
                A[Carga de Certificado de Manufactura en PACT] --> B[Asociar Documento a Partidas Contractuales]
                B --> C{¿El Documento Ampara Múltiples Partidas?}
                C -- Sí --> D[Marcar Opción 'Múltiple' & Asignar Cantidad por Ítem]
                D --> E[Expansión Automática de Registros en Base de Datos]
                C -- No --> F[Registro Directo de Cantidad Respaldada por Ítem]
                E --> G[Verificación de Balance: CM Respaldado vs Cantidad a Pagar]
                F --> G
                G --> H{¿Cantidad CM >= Cantidad Acumulada a Pagar?}
                H -- No --> I[Mapear Alerta Roja: Cantidad de Manufactura Insuficiente]
                I --> J[Bloqueo / Advertencia en Formulario de Pago Mensual]
                H -- Sí --> K[Validar Respaldos Documentales en Verde]
                K --> L[Fin: Partida Habilitada para Cobro 100% Auditado]

                style A fill:#1F4E78,stroke:#38BDF8,color:#fff
                style C fill:#D97706,stroke:#FBBF24,color:#fff
                style I fill:#991B1B,stroke:#F87171,color:#fff
                style K fill:#065F46,stroke:#10B981,color:#fff
            </div>
        </div>

        <!-- 5. MONTHLY PAYMENT -->
        <div class="section-card">
            <h2>5. SECCIÓN: MONTHLY PAYMENT (CERTIFICACIÓN DE PAGO MENSUAL ACT-117C)</h2>
            <div class="description">
                Gestiona la liquidación financiera del periodo. Integra el progreso medido en campo, los respaldos de manufactura, 
                el balance de Material On Site (ACT-117B) y aplica automáticamente la retención contractual.
            </div>
            <div class="mermaid">
            graph TD
                A[Inicio de Periodo de Certificación Mensual] --> B[Conciliar Cantidades Instaladas con Informes ACT-45]
                B --> C[Verificar Validación en Verde de Certificados de Manufactura]
                C --> D[Incorporar Balance de Material On Site ACT-117B]
                D --> E[Ingresar Cantidades a Certificar en el Periodo]
                E --> F[Cálculo Automático del Monto Bruto Acumulado]
                F --> G[Aplicación Automática de Retención Contractual Retainage 5%/10%]
                G --> H[Descuento de Pagos Anteriores & Cálculo de Neto a Pagar]
                H --> I[Revisiones & Firmas: Inspector, Residente & Asesor Legal]
                I --> J[Exportación Oficial de Formulario ACT-117C Excel/PDF]
                J --> K[Fin: Certificación Procesada para Desembolso]

                style A fill:#1F4E78,stroke:#38BDF8,color:#fff
                style F fill:#1E3A8A,stroke:#60A5FA,color:#fff
                style J fill:#065F46,stroke:#10B981,color:#fff
            </div>
        </div>

        <!-- 6. CHANGE ORDER -->
        <div class="section-card">
            <h2>6. SECCIÓN: CHANGE ORDER (ÓRDENES DE CAMBIO CHO / ACT-122 / ACT-123)</h2>
            <div class="description">
                Mecanismo formal y legal para modificar el contrato original. Gestiona partidas extra, ajustes de cantidades, 
                extensiones de plazo y validación presupuestaria federal/estatal.
            </div>
            <div class="mermaid">
            graph TD
                A[Identificación de Necesidad de Cambio en Obra] --> B[Crear Borrador de Orden de Cambio CHO / ACT-122]
                B --> C[Definir Tipo: Ajuste Cantidad / Extra Work / Extensión Tiempo]
                C --> D[Desglosar Costos de Mano de Obra, Equipos & Materiales ACT-123]
                D --> E[Comprobación de Asignación de Fondos Presupuestarios en PACT]
                E --> F{¿El Monto Excede el Techo de Asignación?}
                F -- Sí --> G[Emitir Bloqueo Presupuestario & Requerir Re-asignación]
                F -- No --> H[Completar Check List de Orden de Cambio ACT-124]
                H --> I[Ruta de Firmas Autorizadas: Residente, Director & Agencia]
                I --> J[Aprobación Formal de CHO en PACT]
                J --> K[Inyección Automática de Nuevas Partidas & Ajuste de Calendario]
                K --> L[Fin: Contrato Vigente Actualizado]

                style A fill:#1F4E78,stroke:#38BDF8,color:#fff
                style F fill:#D97706,stroke:#FBBF24,color:#fff
                style G fill:#991B1B,stroke:#F87171,color:#fff
                style K fill:#065F46,stroke:#10B981,color:#fff
            </div>
        </div>
    </div>
    <script>
        mermaid.initialize({
            startOnLoad: true,
            theme: 'dark',
            flowchart: {
                useMaxWidth: true,
                htmlLabels: true,
                curve: 'basis'
            }
        });
    </script>
</body>
</html>
"""
    with open(os.path.join(output_dir, "Diagramas_Flujo_PACT_Interactivas.html"), "w", encoding="utf-8") as f:
        f.write(html_content)

    # 2. DOCUMENTO WORD DOCX DE DIAGRAMAS DE FLUJO
    doc = docx.Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Base style
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Arial'
    style_normal.font.size = Pt(10)
    style_normal.font.color.rgb = RGBColor(0x2B, 0x2B, 0x2B)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run("DIAGRAMAS DE FLUJO DE PROCESOS (FLOWCHARTS)")
    run_title.bold = True
    run_title.font.name = "Arial"
    run_title.font.size = Pt(22)
    run_title.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(14)
    run_sub = p_sub.add_run("Manual de Flujos de Trabajo para las Módulos Principales del Sistema PACT\nPrograma de Administración de Contratos y Control de Obra")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    add_callout(
        doc,
        "SOBRE EL DISEÑO DE PROCESOS PACT",
        "Los diagramas de flujo presentados en este documento representan la arquitectura funcional de negocio del sistema PACT. "
        "Fueron concebidos y diseñados por el Ing. Enrique Saavedra Sada, PE, para normar y automatizar los procesos "
        "de auditoría, entrada de datos de campo, fiscalización laboral, control de manufactura BABA, certificación mensual de pagos "
        "y gestión de órdenes de cambio en proyectos de la Autoridad de Carreteras y Transportación (ACT / PRHTA) y la FHWA.",
        bg_hex="EBF1F5",
        border_hex="1F4E78"
    )

    sections_data = [
        ("1. SECCIÓN: RESUMEN (DASHBOARD & CONSOLIDACIÓN EJECUTIVA)",
         "El módulo de Resumen centraliza los datos maestros del contrato, consolidando el progreso físico de campo (ACT-45), los desembolsos acumulados (ACT-117C) y las modificaciones contractuales (CHO) para alimentar los indicadores ejecutivos (KPIs) y el informe maestro ACT-32.",
         [
             ("Inicio:", "Carga inicial de los datos contractuales y partidas del proyecto."),
             ("Lectura de Datos:", "Procesamiento del Contrato Original y Órdenes de Cambio (CHO) aprobadas."),
             ("Consolidación Acumulada:", "Integración tripartita de Progreso en Campo (ACT-45), Pagos Mensuales (ACT-117C) y Retenciones/MOS (ACT-117B)."),
             ("Cálculo de KPIs:", "Generación instantánea de porcentajes de avance físico, avance financiero, balance por certificar y días ejecutados vs. plazo contractual."),
             ("Salida de Datos:", "Presentación en Tableros Visuales y exportación del Formulario Ejecutivo Maestro ACT-32.")
         ]),

        ("2. SECCIÓN: ENTRADA DE DATOS (INFORMES DIARIOS ACT-45 E INSPECCIÓN ACT-96)",
         "Gobierna la recolección de evidencia primaria en el sitio de construcción. Captura las jornadas laborales, condiciones climáticas, recursos consumidos, cantidades ejecutadas por partida y hallazgos ambientales/seguridad (ACT-96).",
         [
             ("Inicio de Jornada:", "Creación del Informe Diario de Inspección (ACT-45)."),
             ("Registro Ambiental & Recursos:", "Ingreso de temperatura, clima, personal de trabajo y equipo utilizado."),
             ("Selección de Partidas:", "Desplegable estricto de partidas activas del contrato original o CHO aprobadas."),
             ("Ingreso de Cantidades:", "Registro de la producción física diaria por partida."),
             ("Evaluación de Incidencias:", "Si ocurren eventos especiales o inspecciones de agencias externas (EPA, OSHA, ACT, DNER), se abre y vincula el reporte ACT-96."),
             ("Evidencia Fotográfica:", "Anclaje de fotos fehacientes a la fecha y partida trabajada."),
             ("Cierre & Validación:", "Firma digital del Inspector y actualización automática del 'Progreso en Campo' en PACT.")
         ]),

        ("3. SECCIÓN: CUMPLIMIENTO LABORAL (CERTIFIED PAYROLL & DAVIS-BACON)",
         "Módulo fiscalizador de la legislación laboral federal (Davis-Bacon Act) y estatal. Valida las clasificaciones de puestos, tarifas salariales por hora, pago de horas extra y retenciones de nómina de contratistas y subcontratistas.",
         [
             ("Recepción de Nóminas:", "Carga de la nómina semanal certificada del contratista principal y subcontratistas."),
             ("Auditoría de Puestos:", "Verificación de las clasificaciones laborales (carpinteros, operadores, albañiles, etc.)."),
             ("Verificación Salarial:", "Cruce automático de la tarifa pagada por hora contra la escala de salarios mínimos vigentes."),
             ("Control de Horas Extra:", "Auditoría de cálculo de pago de horas extra (Overtime) tras rebasar las horas legales diarias/semanales."),
             ("Validación Jurídica:", "Revisión de la Declaración Jurada adjunta al reporte de nómina."),
             ("Aprobación Laboral:", "Emisión de conformidad en PACT para habilitar el trámite de certificación de pago.")
         ]),

        ("4. SECCIÓN: CERTIFICADOS DE MANUFACTURA (CM & CUMPLIMIENTO BABA)",
         "Aplica la 'Regla de Oro' de PACT. Asegura que ninguna partida de hierro, acero o producto manufacturado BABA sea certifiable para cobro sin contar con el respaldo documental de procedencia y calidad cargado en el sistema.",
         [
             ("Carga Documental:", "Ingreso del Certificado de Manufactura (MTR / BABA) emitido por el fabricante."),
             ("Asociación de Partidas:", "Vinculación directa a la partida o grupo de partidas amparadas."),
             ("Evaluación de Certificado Múltiple:", "Si el certificado cubre varias partidas, se activa la distribución por ítem y PACT expande automáticamente los registros en la base de datos."),
             ("Validación del Balance:", "Comprobación automática: Cantidad Total Certificada CM >= Cantidad Acumulada a Pagar."),
             ("Detección de Incumplimiento:", "Si la cantidad es insuficiente, PACT despliega una alerta roja visual en el módulo de cobro."),
             ("Aprobación:", "Si se cumple el 100% del respaldo, la partida queda habilitada en verde para cobro auditado.")
         ]),

        ("5. SECCIÓN: MONTHLY PAYMENT (CERTIFICACIÓN DE PAGO MENSUAL ACT-117C)",
         "Gestiona la liquidación financiera del periodo. Integra el progreso medido en campo, los respaldos de manufactura, el balance de Material On Site (ACT-117B) y aplica automáticamente la retención contractual.",
         [
             ("Apertura de Periodo:", "Inicio de la Certificación Mensual de Pago (ACT-117C)."),
             ("Conciliación de Campo:", "Carga del acumulado de cantidades trabajadas validadas por los informes ACT-45."),
             ("Verificación de Garantías:", "Confirmación de respaldos de Certificados de Manufactura en verde y adición de Material On Site (ACT-117B)."),
             ("Cálculo Bruto & Retención:", "Determinación del subtotal a pagar y aplicación automática del porcentaje de retención contractual (Retainage 5%/10%)."),
             ("Determinación del Neto:", "Deducción de pagadurías anteriores y establecimiento del monto neto a desembolsar."),
             ("Aprobación & Firma:", "Sello digital del Inspector, Ingeniero Residente y Asesor Legal."),
             ("Exportación Oficial:", "Emisión de reportes oficiales en Excel y PDF.")
         ]),

        ("6. SECCIÓN: CHANGE ORDER (ÓRDENES DE CAMBIO CHO / ACT-122 / ACT-123)",
         "Mecanismo formal y legal para modificar el contrato original. Gestiona partidas extra, ajustes de cantidades, extensiones de plazo y validación presupuestaria federal/estatal.",
         [
             ("Detección del Cambio:", "Identificación de imprevisto en campo, ajuste de diseño o trabajo adicional."),
             ("Elaboración del Borrador:", "Creación de la Orden de Cambio (CHO) mediante formulario ACT-122."),
             ("Desglose de Costos (ACT-123):", "Análisis detallado de mano de obra, equipo, materiales, subcontratos y tiempo adicional."),
             ("Verificación Presupuestaria:", "Validación en PACT de que el monto no sobrepase el techo de asignación de fondos federales/estatales."),
             ("Checklist de Cumplimiento (ACT-124):", "Verificación de endosos ambientales, legales y técnicos."),
             ("Aprobación & Inyección:", "Firma de las partes e inyección automática de nuevas partidas y días de extensión al contrato activo en PACT.")
         ])
    ]

    for title, desc, steps in sections_data:
        p_sec = doc.add_paragraph()
        p_sec.paragraph_format.space_before = Pt(16)
        p_sec.paragraph_format.space_after = Pt(6)
        p_sec.paragraph_format.keep_with_next = True
        run_sec = p_sec.add_run(title)
        run_sec.bold = True
        run_sec.font.name = "Arial"
        run_sec.font.size = Pt(13)
        run_sec.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

        p_d = doc.add_paragraph()
        p_d.paragraph_format.space_before = Pt(0)
        p_d.paragraph_format.space_after = Pt(8)
        p_d.paragraph_format.line_spacing = 1.15
        p_d.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_d = p_d.add_run(desc)
        run_d.font.name = "Arial"
        run_d.font.size = Pt(10)
        run_d.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        p_fl_lbl = doc.add_paragraph()
        p_fl_lbl.paragraph_format.space_before = Pt(4)
        p_fl_lbl.paragraph_format.space_after = Pt(4)
        run_fll = p_fl_lbl.add_run("Secuencia Lógica del Diagrama de Flujo (Pasos del Proceso):")
        run_fll.bold = True
        run_fll.font.name = "Arial"
        run_fll.font.size = Pt(10)
        run_fll.font.color.rgb = RGBColor(0x2F, 0x55, 0x97)

        for step_label, step_text in steps:
            p_st = doc.add_paragraph(style='List Bullet')
            p_st.paragraph_format.space_after = Pt(3)
            run_lbl = p_st.add_run(step_label + " ")
            run_lbl.bold = True
            run_lbl.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
            p_st.add_run(step_text)

    # Section About
    doc.add_page_break()
    p_ab_t = doc.add_paragraph()
    p_ab_t.paragraph_format.space_before = Pt(14)
    p_ab_t.paragraph_format.space_after = Pt(6)
    run_abt = p_ab_t.add_run("SECCIÓN ABOUT - CRÉDITOS Y RECONOCIMIENTO DE AUTORÍA")
    run_abt.bold = True
    run_abt.font.name = "Arial"
    run_abt.font.size = Pt(15)
    run_abt.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

    add_callout(
        doc,
        "SISTEMA PACT - AUTORÍA Y PROPIEDAD INTELECTUAL",
        "El diseño de flujos de proceso, arquitectura de datos y diagramas funcionales del programa PACT "
        "ha sido desarrollado en su totalidad por el Ing. Enrique Saavedra Sada, PE.\n\n"
        "Esta documentación técnica de flujos de trabajo (Flowcharts) sirve como estándar oficial de operación "
        "para la administración, inspección y auditoría de proyectos de construcción bajo el sistema PACT.\n\n"
        "Diseñador del Software y Procesos: Ing. Enrique Saavedra Sada, PE\n"
        "Fecha de Publicación: Julio 2026\n"
        "Ubicación de Archivos: C:\\Users\\Enrique Saavedra\\Documents\\PROGRAMAS AI\\Programa ACT\\Documentos\\Flowcharts",
        bg_hex="F2F4F7",
        border_hex="1F4E78"
    )

    doc_out_path = os.path.join(output_dir, "Flowcharts_Secciones_PACT.docx")
    doc.save(doc_out_path)
    print(f"Documento de Flowcharts guardado con éxito en: {doc_out_path}")

if __name__ == "__main__":
    create_flowcharts()
