import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_callout(doc, title, text, bg_hex="F2F4F7", border_hex="1F4E78"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    # Left border
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24') # 3pt
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), border_hex)
    tcBorders.append(left)
    
    for side in ['top', 'bottom', 'right']:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:val'), 'none')
        tcBorders.append(node)
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run_t = p.add_run(f"📌 {title}\n")
    run_t.bold = True
    run_t.font.name = "Arial"
    run_t.font.size = Pt(11)
    run_t.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    
    run_b = p.add_run(text)
    run_b.font.name = "Arial"
    run_b.font.size = Pt(10)
    run_b.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # Empty paragraph after table for spacing
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(0)
    p_space.paragraph_format.space_after = Pt(6)

def create_faq_document():
    doc = docx.Document()

    # Set Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styles setup
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x2B, 0x2B, 0x2B)

    # Document Header / Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run("GUÍA DE PREGUNTAS FRECUENTES (FAQ)")
    run_title.bold = True
    run_title.font.name = "Arial"
    run_title.font.size = Pt(22)
    run_title.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(12)
    run_sub = p_sub.add_run("Sistema PACT – Programa de Administración de Contratos y Control de Obra\n40 Respuestas Clave para Profesionales de la Construcción e Ingeniería")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    # Designer Banner Box
    add_callout(
        doc,
        "SOBRE EL SISTEMA PACT",
        "Este documento ha sido diseñado como guía práctica de aprendizaje y referencia técnica para ingenieros, inspectores y administradores de proyectos junior. El sistema PACT fue concebido, diseñado y desarrollado por el Ing. Enrique Saavedra Sada, PE, con el objetivo de garantizar la integridad conceptual, presupuestaria y legal en la administración de contratos de obras públicas y privadas.",
        bg_hex="EBF1F5",
        border_hex="1F4E78"
    )

    faqs = [
        # CATEGORIA I
        ("CATEGORÍA I: CONCEPTOS FUNDAMENTALES Y ESTRUCTURA GENERAL DE PACT", [
            ("1. ¿Qué es el programa PACT y cuál es su objetivo principal?",
             "PACT (Programa de Administración de Contratos) es una solución especializada de software para el control técnico, presupuestario y legal de proyectos de construcción. Su objetivo principal es automatizar la certificación de pagos, el control de informes diarios de inspección, el seguimiento de órdenes de cambio y el respaldo con certificados de manufactura, garantizando que cada centavo certificado esté debidamente auditado y respaldado."),
            
            ("2. ¿A quién está dirigida esta guía de preguntas frecuentes (FAQ)?",
             "Esta guía está orientada a profesionales con poca experiencia o de nivel junior (ingenieros residentes, inspectores de campo, administradores de contratos y auxiliares de contabilidad) que necesitan dominar los detalles operativos y las reglas de negocio del sistema PACT para evitar errores de auditoría y garantizar cobros fluidos."),
            
            ("3. ¿Cómo calcula PACT la duración de los días en el calendario del proyecto?",
             "Siguiendo las reglas de auditoría de PACT, para que un periodo o jornada sea considerado legalmente como un día dentro de los cálculos contractuales y de plazo ejecutado, debe contabilizarse como un periodo completo de 24 horas. Esto garantiza precisión al medir extensiones de tiempo y demoras."),
            
            ("4. ¿Qué es el Audit Log (Registro de Auditoría) en PACT y por qué no se puede desactivar?",
             "El Audit Log es un registro inalterable donde PACT almacena de forma automática cualquier acción crítica (creación, edición o eliminación de informes diarios, certificaciones o órdenes de cambio), registrando la fecha, hora exacta y usuario. No se puede desactivar porque garantiza la trazabilidad total requerida en auditorías estatales y federales."),
            
            ("5. ¿Qué sucede internamente si un usuario elimina un Informe Diario o una Certificación de Pago?",
             "Cuando un usuario elimina un registro, PACT actualiza inmediatamente el Audit Log y recalcula de forma automática todas las cantidades acumuladas del proyecto para evitar discrepancias. Si se elimina un informe diario, el progreso en campo de las partidas involucradas retrocede automáticamente al estado previo."),
            
            ("6. ¿Dónde puedo verificar la autoría y versión oficial del programa PACT?",
             "En la sección “About” (Acerca de) dentro de la interfaz del programa PACT. Allí se indica de forma explícita la versión activa del sistema (ej. v3.28.0503) y se otorga el reconocimiento al diseñador y creador del software: Ing. Enrique Saavedra Sada, PE.")
        ]),

        # CATEGORIA II
        ("CATEGORÍA II: GESTIÓN DE CONTRATOS, PRESUPUESTO Y PARTIDAS (ITEMS)", [
            ("7. ¿Cómo se compone la estructura de partidas (Items) en PACT?",
             "El contrato en PACT se compone de las partidas del Contrato Original (con sus descripciones, unidades de medida, cantidades contratadas y precios unitarios) más las partidas adicionadas o modificadas mediante Órdenes de Cambio (CHO / ACT-123)."),
            
            ("8. ¿Se pueden modificar las partidas o precios del Contrato Original directamente en la Certificación de Pago?",
             "No. PACT prohíbe la alteración directa de descripciones, unidades o precios unitarios desde la pantalla de certificación. Las únicas vías legales para alterar el presupuesto o las partidas son la aprobación formal de una Orden de Cambio (CHO / ACT-123)."),
            
            ("9. ¿Cuál es la diferencia entre certificar partidas a Precio Unitario versus Suma Global (Lump Sum)?",
             "Las partidas a Precio Unitario se certifican ingresando la cantidad física medida e instalada en campo (ej. metros cúbicos de hormigón). Las partidas a Suma Global (Lump Sum) se certifican por porcentaje de avance acumulado o fracciones estipuladas según el desglose de valores aprobado."),
            
            ("10. ¿Cómo distingue PACT entre el Presupuesto Original y el Presupuesto Vigente?",
             "El Presupuesto Original representa el monto adjudicado al firmar el contrato. El Presupuesto Vigente es la suma matemática del Presupuesto Original más o menos el impacto neto de todas las Órdenes de Cambio (CHO) debidamente aprobadas."),
            
            ("11. ¿Cómo valida PACT la disponibilidad de fondos presupuestarios?",
             "PACT efectúa una comprobación en tiempo real: valida que la suma acumulada de partidas y solicitudes de pago no supere el techo de asignación de fondos (federales o estatales) autorizados para el proyecto. Si se excede, el sistema emite un bloqueo o alerta presupuestaria."),
            
            ("12. ¿Qué ocurre si intento asociar un trabajo a una partida inexistente?",
             "El sistema no permitirá guardar el registro. En los formularios de campo (como el ACT-45), PACT despliega un menú desplegable estricto que solo contiene partidas activas e inyectadas desde el contrato original o CHO aprobadas.")
        ]),

        # CATEGORIA III
        ("CATEGORÍA III: INFORMES DIARIOS DE CAMPO (ACT-45) E INSPECCIÓN (ACT-96)", [
            ("13. ¿Qué es el Informe Diario ACT-45 y cuál es su importancia operativa?",
             "El ACT-45 es el documento primario de campo donde el inspector registra las actividades diarias, mano de obra, equipo utilizado, condiciones climáticas y cantidades instaladas por partida. Es la fuente fundamental que alimenta el progreso acumulado en PACT."),
            
            ("14. ¿Cómo alimentan los informes ACT-45 a las Certificaciones de Pago Mensuales?",
             "Cada vez que un inspector registra cantidades ejecutadas en un ACT-45, PACT actualiza en tiempo real el 'Progreso en Campo'. Al preparar la Certificación de Pago Mensual, el sistema sugiere las cantidades acumuladas trabajadas, facilitando la conciliación entre el inspector y el contratista."),
            
            ("15. ¿Qué se debe hacer si se detecta un error de digitación en un ACT-45 de semanas pasadas?",
             "El usuario autorizado debe editar el informe ACT-45 correspondiente. Al guardar los cambios, PACT recalculará automáticamente los acumulados de las partidas afectadas y dejará constancia de la modificación en el Audit Log."),
            
            ("16. ¿Qué es el formulario ACT-96 y en qué se diferencia del ACT-45?",
             "El ACT-96 es el Informe de Inspección Específica y Gestión de Riesgos. Mientras que el ACT-45 detalla la producción y recursos diarios, el ACT-96 documenta eventos ambientales, de seguridad, pruebas de laboratorio o visitas de agencias de fiscalización."),
            
            ("17. ¿Cómo gestiona PACT las inspecciones de agencias externas (EPA, OSHA, ACT, DNER)?",
             "Las visitas e hallazgos de agencias externas se registran en el módulo ACT-96. Si se marca una falla de cumplimiento grave o incumplimiento ambiental, PACT permite vincular dicho hallazgo para bloquear provisionalmente la certificación de las partidas asociadas hasta que se subsane."),
            
            ("18. ¿Cómo se integran y exportan las fotos en los Informes Diarios ACT-45?",
             "Las fotografías capturadas en campo se anclan a la fecha y a las partidas específicas del ACT-45. Al exportar el reporte a Excel o PDF, PACT compila y genera automáticamente la hoja de anexo fotográfico estructurado.")
        ]),

        # CATEGORIA IV
        ("CATEGORÍA IV: CERTIFICACIÓN DE PAGOS MENSUALES, ACUMULADOS Y EXCEDENTES", [
            ("19. ¿Cuál es el flujo estándar para generar una Certificación de Pago Mensual?",
             "1) Conciliar cantidades ejecutadas en campo mediante los informes ACT-45. 2) Verificar respaldos de Certificados de Manufactura. 3) Ingresar las cantidades a pagar en el periodo. 4) Verificar retenciones aplicables. 5) Someter el reporte a revisión del Ingeniero Residente."),
            
            ("20. ¿Cómo calcula PACT la 'Cantidad Acumulada Anterior' y 'Cantidad Acumulada Actual'?",
             "La 'Cantidad Acumulada Anterior' es la suma de todas las cantidades certificadas en los periodos de pago 1 hasta (N-1). La 'Cantidad Acumulada Actual' es la suma de la Cantidad Acumulada Anterior más la 'Cantidad a Pagar en este Periodo'."),
            
            ("21. ¿Qué es la Retención (Retainage) en PACT y cómo se aplica?",
             "Es un porcentaje (habitualmente 5% o 10%) que se descuenta del valor bruto a pagar en cada certificación para garantizar la fiel ejecución de la obra. PACT calcula este valor automáticamente sobre el total de trabajo certificado de cada periodo."),
            
            ("22. ¿Qué ocurre en PACT cuando la cantidad a pagar excede la cantidad contratada (Excedente de Certificación)?",
             "PACT detecta que la Cantidad Acumulada Actual supera la Cantidad Contratada de la partida. Si el excedente no está amparado por una Orden de Cambio (CHO), el sistema genera una advertencia de sobre-certificación para evitar pagos improcedentes."),
            
            ("23. ¿Cómo maneja PACT la certificación de Pagos Iniciales o Anticipos de Materiales (Stored Materials)?",
             "PACT cuenta con módulos específicos para certificar materiales acopiados en obra o manufacturados que aún no han sido instalados. Requiere que se adjunte la factura y el certificado de propiedad antes de conceder el pago parcial."),
            
            ("24. ¿Es posible modificar una Certificación de Pago que ya ha sido aprobada y sellada?",
             "Una vez que una certificación pasa a estado 'Aprobado/Finalizado', queda bloqueada para edición directa. Para corregir un error, se debe reabrir mediante un permiso de rol superior o realizar el ajuste compensatorio en la certificación del periodo siguiente."),
            
            ("25. ¿Qué checklist mínimo debe revisar un ingeniero junior antes de enviar a firma la Certificación de Pago?",
             "✓ Verificación de cantidades en informes diarios ACT-45.\n✓ Cobertura del 100% de la cantidad por Certificados de Manufactura.\n✓ Verificación de que no existan partidas bloqueadas por incumplimiento en ACT-96.\n✓ Correcta aplicación de la retención contractual.")
        ]),

        # CATEGORIA V
        ("CATEGORÍA V: CERTIFICADOS DE MANUFACTURA Y EVIDENCIA DOCUMENTAL", [
            ("26. ¿En qué consiste la 'Regla de Oro' de PACT respecto a los Certificados de Manufactura?",
             "La 'Regla de Oro' establece que **no se puede certificar el pago de una partida si no existe evidencia documental (Certificado de Manufactura) que respalde la cantidad instalada**."),
            
            ("27. ¿Qué alerta visual muestra PACT si no hay suficiente Certificado de Manufactura?",
             "Si la cantidad a pagar supera lo respaldado por los certificados subidos, PACT muestra una alerta visual destacada en rojo: *'Advertencia: La cantidad certificada de manufactura para esta partida es insuficiente.'*"),
            
            ("28. ¿Qué es un Certificado de Manufactura Múltiple?",
             "Es un documento (por ejemplo, una prueba de resistencia o calidad de la planta siderúrgica o de concreto) que ampara simultáneamente varias partidas del contrato en un solo certificado."),
            
            ("29. ¿Cómo procesa PACT un Certificado de Manufactura Múltiple de forma automática?",
             "Al marcar un certificado como 'Múltiple', PACT permite seleccionar todas las partidas amparadas e ingresar la cantidad específica asignada a cada una. Al guardar, el sistema 'expande' automáticamente ese registro único en entradas individuales por partida en la base de datos."),
            
            ("30. ¿Por qué el proceso de expansión automática de certificados facilita la auditoría?",
             "Porque permite rastrear de forma individual e instantánea cuántos Certificados de Manufactura respaldan a cada partida específica sin necesidad de revisar manualmente documentos multipágina repetidos.")
        ]),

        # CATEGORIA VI
        ("CATEGORÍA VI: ÓRDENES DE CAMBIO (CHO / ACT-123) Y MODIFICACIONES PRESUPUESTARIAS", [
            ("31. ¿Qué es una Orden de Cambio (CHO / ACT-123) en PACT?",
             "Es el instrumento legal formal para modificar las condiciones originales del contrato, ya sea alterando cantidades, agregando partidas extras (Extra Work), suprimiendo partidas o ajustando el costo y plazo de la obra."),
            
            ("32. ¿Cómo impacta una CHO aprobada al módulo de Certificación de Pagos?",
             "Una vez que una CHO/ACT-123 pasa a estado 'Aprobado', PACT inyecta automáticamente las nuevas partidas en la tabla de certificación o actualiza las cantidades autorizadas de las partidas existentes."),
            
            ("33. ¿Se pueden certificar pagos sobre una Orden de Cambio (CHO) en borrador?",
             "No. PACT exige que la CHO esté legalmente aprobada en el sistema para habilitar la certificación y cobro de sus partidas de trabajo."),
            
            ("34. ¿Cómo afecta una extensión de tiempo otorgada mediante CHO al calendario de PACT?",
             "Al aprobarse días de extensión en una CHO, PACT recalcula automáticamente la Fecha Substancial de Terminación y la Fecha de Terminación Final del contrato, ajustando los indicadores de demora."),
            
            ("35. ¿Qué documentación mínima exige PACT para sustentar una Orden de Cambio?",
             "Justificación técnica, desglose de costos de mano de obra y equipos, cotizaciones, planos/esquemas aprobados y el acuerdo de tiempo firmado por las partes autorizadas.")
        ]),

        # CATEGORIA VII
        ("CATEGORÍA VII: FIRMAS AUTORIZADAS, TRAZABILIDAD, AUDITORÍA Y EXPORTACIÓN", [
            ("36. ¿Cómo se estructura la jerarquía de firmas en PACT?",
             "• **Inspector de Campo:** Sella reportes diarios (ACT-45) e inspecciones (ACT-96).\n• **Ingeniero Residente:** Valida y aprueba la Certificación de Pago Mensual.\n• **Director de Área / Asesor Legal:** Firman reportes maestros consolidados (como el ACT-32)."),
            
            ("37. ¿Dónde se configuran los nombres y títulos oficiales que aparecen en los reportes?",
             "Se configuran en la sección **Firmas Autorizadas** de la configuración del proyecto. PACT inyecta automáticamente estos nombres, puestos y licencias profesionales en los archivos Excel y PDF generados."),
            
            ("38. ¿Qué es el reporte maestro ACT-32 y cuál es su función?",
             "El ACT-32 es el resumen ejecutivo consolidado del estado financiero y físico del contrato. Muestra el monto contratado, total de CHO aprobadas, valor certificado a la fecha, retenciones acumuladas y balance pendiente de pago."),
            
            ("39. ¿Cuáles son las mejores prácticas de respaldo de datos recomendadas en PACT?",
             "Se recomienda realizar exportaciones semanales de la base de datos de PACT y respaldar las carpetas de adjuntos (fotos y certificados de manufactura) en almacenamiento seguro o nube."),
            
            ("40. ¿Qué contiene la sección 'About' de PACT y por qué debe preservarse?",
             "La sección 'About' acredita formalmente el origen del software, indicando que ha sido concebido y diseñado por el **Ing. Enrique Saavedra Sada, PE**. Debe preservarse para respetar los derechos de propiedad intelectual y validar la autenticidad del programa.")
        ])
    ]

    for category_title, qa_list in faqs:
        # Category Title
        p_cat = doc.add_paragraph()
        p_cat.paragraph_format.space_before = Pt(18)
        p_cat.paragraph_format.space_after = Pt(8)
        p_cat.paragraph_format.keep_with_next = True
        run_cat = p_cat.add_run(category_title)
        run_cat.bold = True
        run_cat.font.name = "Arial"
        run_cat.font.size = Pt(13)
        run_cat.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

        # Draw a subtle line under category
        for q_title, q_ans in qa_list:
            p_q = doc.add_paragraph()
            p_q.paragraph_format.space_before = Pt(8)
            p_q.paragraph_format.space_after = Pt(2)
            p_q.paragraph_format.keep_with_next = True
            
            run_qt = p_q.add_run(q_title)
            run_qt.bold = True
            run_qt.font.name = "Arial"
            run_qt.font.size = Pt(11)
            run_qt.font.color.rgb = RGBColor(0x2F, 0x55, 0x97)

            p_a = doc.add_paragraph()
            p_a.paragraph_format.space_before = Pt(0)
            p_a.paragraph_format.space_after = Pt(6)
            p_a.paragraph_format.line_spacing = 1.15
            p_a.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            run_at = p_a.add_run(q_ans)
            run_at.font.name = "Arial"
            run_at.font.size = Pt(10)
            run_at.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Final About Section
    doc.add_page_break()
    p_ab_title = doc.add_paragraph()
    p_ab_title.paragraph_format.space_before = Pt(12)
    p_ab_title.paragraph_format.space_after = Pt(6)
    run_abt = p_ab_title.add_run("SECCIÓN ABOUT - CRÉDITOS DEL SOFTWARE")
    run_abt.bold = True
    run_abt.font.name = "Arial"
    run_abt.font.size = Pt(16)
    run_abt.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

    add_callout(
        doc,
        "SISTEMA PACT - ACERCA DEL DISEÑADOR",
        "El programa PACT (Programa de Administración de Contratos y Control de Obra) ha sido diseñado e implementado por el Ing. Enrique Saavedra Sada, PE.\n\n"
        "Esta herramienta integra las mejores prácticas de ingeniería de construcción, control presupuestario y cumplimiento legal para proyectos de alta exigencia técnica.\n\n"
        "Diseñador del Software: Ing. Enrique Saavedra Sada, PE\n"
        "Versión del Manual / Guía FAQ: 1.0 (Julio 2026)\n"
        "Ubicación de Documentos: C:\\Users\\Enrique Saavedra\\Documents\\PROGRAMAS AI\\Programa ACT\\Documentos",
        bg_hex="F2F4F7",
        border_hex="1F4E78"
    )

    out_dir = r"C:\Users\Enrique Saavedra\Documents\PROGRAMAS AI\Programa ACT\Documentos"
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, "Preguntas_Frecuentes_PACT.docx")
    doc.save(out_path)
    print(f"Documento guardado con éxito en: {out_path}")

if __name__ == "__main__":
    create_faq_document()
