import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="B0C4DE", sz="4", val="single"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>\n'
            f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
            f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
            f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
            f'  <w:insideV w:val="none"/>\n'
            f'  <w:left w:val="none"/>\n'
            f'  <w:right w:val="none"/>\n'
            f'</w:tblBorders>'
        )
        tblPr[0].append(borders)

def add_callout(doc, title, text, bg_hex="F0F4F8", border_hex="003366"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>\n'
        f'  <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_hex}"/>\n'
        f'  <w:top w:val="none"/>\n'
        f'  <w:right w:val="none"/>\n'
        f'  <w:bottom w:val="none"/>\n'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run_t = p.add_run(f"📌 {title}\n")
    run_t.bold = True
    run_t.font.name = "Arial"
    run_t.font.size = Pt(11)
    run_t.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    
    run_b = p.add_run(text)
    run_b.font.name = "Arial"
    run_b.font.size = Pt(10)
    run_b.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(0)
    p_space.paragraph_format.space_after = Pt(6)

def add_toc_field(paragraph):
    run = paragraph.add_run()
    fldChar1 = parse_xml(r'<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w'))
    instrText = parse_xml(r'<w:instrText %s xml:space="preserve"> TOC \o "1-3" \h \z \u </w:instrText>' % nsdecls('w'))
    fldChar2 = parse_xml(r'<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w'))
    fldChar3 = parse_xml(r'<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w'))
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def generate_baba_document():
    doc = docx.Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Base Typography
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Arial'
    style_normal.font.size = Pt(10)
    style_normal.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    # Header / Title Block
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(6)
    run_title = p_title.add_run("ANÁLISIS DE ÍTEMS QUE REQUIEREN CERTIFICADO DE MANUFACTURA BAJO BUILD AMERICA, BUY AMERICA (BABA)")
    run_title.bold = True
    run_title.font.name = "Arial"
    run_title.font.size = Pt(20)
    run_title.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(18)
    run_sub = p_sub.add_run("Con Especificaciones Estándar ACT/PRHTA de Puerto Rico (Lista Histórica de Partidas)\nPara Proyectos de Carreteras y Puentes Financiados por la FHWA")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(11.5)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Callout Box Metadata
    add_callout(
        doc,
        "DATOS DEL DOCUMENTO Y REFERENCIA TÉCNICA ACT",
        "Documento Técnico de Cumplimiento Normativo Federal BABA / FHWA\n"
        "Matriz Vinculada al Libro de Especificaciones Estándar de Carreteras (Spec History ACT)\n"
        "Diseñado y Preparado por: Ing. Enrique Saavedra Sada, PE\n"
        "Programa / Sistema: PACT (Programa de Administración de Contratos)\n"
        "Dirigido a: Administradores de Contratos, Ingenieros Residentes, Inspectores de Campo y Oficiales de Cumplimiento BABA/FHWA",
        bg_hex="F4F6F9",
        border_hex="003366"
    )

    # TABLA DE CONTENIDO AUTOMÁTICA
    p_toc_heading = doc.add_paragraph()
    p_toc_heading.paragraph_format.space_before = Pt(18)
    p_toc_heading.paragraph_format.space_after = Pt(6)
    run_toch = p_toc_heading.add_run("TABLA DE CONTENIDO")
    run_toch.bold = True
    run_toch.font.name = "Arial"
    run_toch.font.size = Pt(14)
    run_toch.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    p_toc = doc.add_paragraph()
    p_toc.paragraph_format.space_after = Pt(18)
    add_toc_field(p_toc)

    doc.add_page_break()

    # --- SECCIÓN 1 ---
    h1 = doc.add_heading(level=1)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(8)
    run_h1 = h1.add_run("1. INTRODUCCIÓN Y MARCO NORMATIVO BABA / FHWA CON ESPECIFICACIONES ACT")
    run_h1.bold = True
    run_h1.font.name = "Arial"
    run_h1.font.size = Pt(15)
    run_h1.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    p_intro = doc.add_paragraph()
    p_intro.paragraph_format.line_spacing = 1.15
    p_intro.paragraph_format.space_after = Pt(8)
    p_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_intro.add_run(
        "El Ley de Inversión en Infraestructura y Empleos (IIJA / Ley Bipartidista de Infraestructura, P.L. 117-58), "
        "estableció la Ley Build America, Buy America (BABA) bajo la Ley Orgánica de Infraestructura. Esta disposición amplía "
        "y refuerza los requisitos históricos de 'Buy America' de la Administración Federal de Carreteras (FHWA, 23 CFR 635.410), "
        "exigiendo que todas las subvenciones y fondos de asistencia federal otorgados a proyectos de infraestructura de carreteras "
        "y puentes en Estados Unidos y Puerto Rico utilicen productos fabricados nacionalmente."
    )

    p_intro2 = doc.add_paragraph()
    p_intro2.paragraph_format.line_spacing = 1.15
    p_intro2.paragraph_format.space_after = Pt(8)
    p_intro2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_intro2.add_run(
        "Para garantizar la operabilidad práctica en Puerto Rico, este documento vincula cada ítem evaluado con su correspondiente "
        "**Código de Especificación Estándar (Spec Code)** de la Autoridad de Carreteras y Transportación (ACT / PRHTA), extraído "
        "de la base de datos oficial del historial de partidas (Items History)."
    )

    # 1.1 Criterios de Evaluación
    h1_1 = doc.add_heading(level=2)
    h1_1.paragraph_format.space_before = Pt(12)
    h1_1.paragraph_format.space_after = Pt(6)
    run_h1_1 = h1_1.add_run("1.1 Criterios para Determinar si un Ítem Requiere Certificado de Manufactura")
    run_h1_1.bold = True
    run_h1_1.font.name = "Arial"
    run_h1_1.font.size = Pt(12)
    run_h1_1.font.color.rgb = RGBColor(0x2F, 0x55, 0x97)

    criterios = [
        ("1. Incorporación Permanente:", "El artículo o producto se instala para permanecer de forma definitiva en la obra terminada."),
        ("2. Categoría Cubierta por BABA:", "El artículo pertenece a una de las categorías reguladas: Hierro y Acero, Productos Manufacturados, Materiales de Construcción o Productos Combinados."),
        ("3. Financiamiento Federal FHWA:", "El producto es adquirido mediante fondos federales o subvenciones administradas por FHWA / USDOT."),
        ("4. Trazabilidad del Fabricante:", "El origen nacional estadounidense debe demostrarse documentalmente mediante certificación formal emitida por el fabricante.")
    ]
    for tit, desc in criterios:
        p_item = doc.add_paragraph(style='List Bullet')
        p_item.paragraph_format.space_after = Pt(4)
        run_b = p_item.add_run(tit + " ")
        run_b.bold = True
        run_b.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
        p_item.add_run(desc)

    # 1.2 Requisitos del Certificado
    h1_2 = doc.add_heading(level=2)
    h1_2.paragraph_format.space_before = Pt(12)
    h1_2.paragraph_format.space_after = Pt(6)
    run_h1_2 = h1_2.add_run("1.2 Contenido Mínimo Exigido en el Certificado de Manufactura")
    run_h1_2.bold = True
    run_h1_2.font.name = "Arial"
    run_h1_2.font.size = Pt(12)
    run_h1_2.font.color.rgb = RGBColor(0x2F, 0x55, 0x97)

    req_list = [
        "Que el producto fue manufacturado en los Estados Unidos de América.",
        "Que cumple estrictamente con los estándares aplicables de Buy America (23 CFR 635.410) y BABA (2 CFR 184).",
        "Que **todos los procesos de manufactura**, desde la fusión inicial (melting) hasta la aplicación de recubrimientos (coating, galvanizado, pintura epóxica), ocurrieron en territorio estadounidense (aplicable a hierro y acero).",
        "Que para materiales de construcción y productos manufacturados, el costo de componentes producidos en EE.UU. excede el 55% del costo total de componentes.",
        "Identificación del número de lote, orden de compra, especificación estándar ACT aplicable (ej. Spec 602, Spec 605, Spec 610, Spec 624, Spec 625) y partida contractual asociada."
    ]
    for r in req_list:
        p_r = doc.add_paragraph(style='List Bullet')
        p_r.paragraph_format.space_after = Pt(3)
        p_r.add_run(r)

    # --- SECCIÓN 2 ---
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    h2 = doc.add_heading(level=1)
    h2.paragraph_format.space_before = Pt(16)
    h2.paragraph_format.space_after = Pt(8)
    run_h2 = h2.add_run("2. SECCIÓN A – ÍTEMS QUE EXPRESAMENTE REQUIEREN CERTIFICADO CON ESPECIFICACIÓN ACT")
    run_h2.bold = True
    run_h2.font.name = "Arial"
    run_h2.font.size = Pt(14)
    run_h2.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    # TABLA SECCIÓN A CON COLUMNA ESPECIFICACIÓN ACT
    table_a = doc.add_table(rows=1, cols=5)
    table_a.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table_a, color="B0C4DE", sz="6")

    hdr_cells = table_a.rows[0].cells
    hdr_titles = ["Ítem", "Especificación ACT", "Categoría BABA", "Razón por la cual requiere certificado", "Tipo de Certificado Requerido"]
    col_widths_a = [Inches(1.2), Inches(1.1), Inches(1.2), Inches(1.8), Inches(1.2)]

    for i, title in enumerate(hdr_titles):
        hdr_cells[i].width = col_widths_a[i]
        set_cell_background(hdr_cells[i], "003366")
        set_cell_margins(hdr_cells[i], top=140, bottom=140, left=100, right=100)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(title)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    data_a = [
        ("Rebar (Varilla de Refuerzo)", "Spec 602 / 889\n(Reinforcing Steel)", "Hierro y Acero", "Componente estructural permanente de acero sujeto a norma estricta 23 CFR 635.410. Fusión y rolado 100% doméstico.", "Certificado de Mill (MTR) + Certificación Buy America 100% US Melt"),
        ("Acero Estructural (Vigas, Diafragmas)", "Spec 605 / 843 / 853\n(Structural Steel)", "Hierro y Acero", "Elemento sustentante primario en puentes. Requiere trazabilidad total del colado, fabricación y galvanizado.", "MTR de Planta + Certificado de Fabricación y Galvanizado Doméstico"),
        ("Guardrails (Barreras Metálicas)", "Spec 610 / 834\n(Guardrail Systems)", "Hierro y Acero / Prod. Combinado", "Lámina de acero corrugado y postes metálicos con galvanizado por inmersión en caliente (ASTM A123).", "Certificado de Chapa/Postes + Certificado de Galvanizado ASTM A123"),
        ("Postes de Señales de Tránsito", "Spec 622 / 904\n(Sign Posts & Supports)", "Hierro y Acero", "Estructuras tubulares fijadas permanentemente. El acero y recubrimiento anticorrosivo deben ser nacionales.", "Certificado de Fabricante de Acero + Certificación de Galvanizado"),
        ("Señales Permanentes", "Spec 622\n(Traffic Signs)", "Producto Manufacturado", "Láminas de aluminio con sustrato reflectivo permanente. Cumplen con regla BABA de productos manufacturados.", "Certificado de Cumplimiento BABA del Fabricante de la Señal y Película"),
        ("Postes de Alumbrado Público", "Spec 624 / 970\n(Highway Lighting Poles)", "Hierro y Acero / Prod. Manufacturado", "Estructuras tubulares de acero o aluminio sobre bases de hormigón. El metal y su ensamble deben ser domésticos.", "Certificado de Fabricante de Poste Metálico y Ensamble Nacional"),
        ("Tubería de Hierro Dúctil", "Spec 603 / 604\n(Ductile Iron Pipe)", "Hierro y Acero", "Tubería subterránea incorporada de forma permanente. Clasificada bajo la exigencia estricta de acero/hierro BABA.", "Certificado de Fundición y Manufactura 100% US Melt (Ductile Iron MTR)"),
        ("Dynamic Message Signs (DMS)", "Spec 622 / 625\n(Dynamic Message Signs)", "Producto Manufacturado / ITS", "Equipo electrónico complejo incorporado a la red vial. Ensamblado y componentes deben cumplir BABA (>55%).", "Certificado BABA de Producto Manufacturado + Desglose de Componentes"),
        ("Cámaras ITS y Sensores", "Spec 625 / 833\n(ITS Surveillance / Camera)", "Producto Manufacturado", "Dispositivos de monitoreo vial permanente integrados a la infraestructura pública de transportación.", "Certificación BABA del Fabricante del Sistema ITS"),
        ("Componentes de Puentes", "Spec 606 / 939 / 979\n(Expansion Joints/Bearings)", "Producto Combinado / Mat. Construcción", "Soportes elastoméricos con placas de acero internas y juntas de expansión metálicas permanentes.", "Certificado de Ensamble y Vulcanizado con Acero Certificado US Melt")
    ]

    for idx, row in enumerate(data_a):
        row_cells = table_a.add_row().cells
        bg_color = "F9FAFC" if idx % 2 == 1 else "FFFFFF"
        for i, val in enumerate(row):
            row_cells[i].width = col_widths_a[i]
            set_cell_background(row_cells[i], bg_color)
            set_cell_margins(row_cells[i], top=90, bottom=90, left=80, right=80)
            p = row_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(val)
            run.font.name = "Arial"
            run.font.size = Pt(8.5)
            if i == 1:
                run.bold = True
                run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
            else:
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # --- SECCIÓN 3 ---
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    h3 = doc.add_heading(level=1)
    h3.paragraph_format.space_before = Pt(16)
    h3.paragraph_format.space_after = Pt(8)
    run_h3 = h3.add_run("3. SECCIÓN B – ÍTEMS QUE DEBERÍAN REQUERIR CERTIFICADO CON ESPECIFICACIÓN ACT")
    run_h3.bold = True
    run_h3.font.name = "Arial"
    run_h3.font.size = Pt(14)
    run_h3.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    # TABLA SECCIÓN B CON COLUMNA ESPECIFICACIÓN ACT
    table_b = doc.add_table(rows=1, cols=5)
    table_b.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table_b, color="B0C4DE", sz="6")

    hdr_cells_b = table_b.rows[0].cells
    hdr_titles_b = ["Ítem / Componente", "Especificación ACT", "Categoría BABA", "Justificación Técnica y Alcance", "Probabilidad"]
    col_widths_b = [Inches(1.3), Inches(1.1), Inches(1.1), Inches(2.2), Inches(0.8)]

    for i, title in enumerate(hdr_titles_b):
        hdr_cells_b[i].width = col_widths_b[i]
        set_cell_background(hdr_cells_b[i], "003366")
        set_cell_margins(hdr_cells_b[i], top=140, bottom=140, left=100, right=100)
        p = hdr_cells_b[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(title)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    data_b = [
        ("Barreras de Seguridad de Hormigón", "Spec 601 / 901 / 906\n(Concrete Barrier)", "Producto Combinado", "Barreras New Jersey/F-Shape con varillas de refuerzo permanentes. El acero debe ser 100% US Melt.", "Alto"),
        ("Atenuadores de Impacto", "Spec 610 / 900 / 920\n(Impact Attenuators)", "Producto Manufacturado / Hierro", "Sistemas mecánicos de absorción de energía fijados al pavimento. Requieren certificado del fabricante.", "Alto"),
        ("Postes Galvanizados de Alumbrado", "Spec 624 / 970\n(Galvanized Lighting Poles)", "Hierro y Acero", "Estructuras de soporte a intemperie. La fabricación del tubo y galvanizado deben ser en EE.UU.", "Alto"),
        ("Sistemas de Alumbrado y Luminarias LED", "Spec 624 / 825 / 951 / 988\n(Luminaires / LED)", "Producto Manufacturado", "Luminarias de alta eficiencia fijadas permanentemente. Componentes domésticos deben superar el 55%.", "Alto"),
        ("Semáforos y Signal Heads", "Spec 625 / 956\n(Traffic Signals)", "Producto Manufacturado", "Cabezales de aluminio/polímero e iluminación LED incorporados permanentemente al sistema de tráfico.", "Alto"),
        ("Gabinetes Eléctricos / Tráfico", "Spec 624 / 625\n(Electrical Cabinets)", "Hierro y Acero / Prod. Manuf.", "Cajas NEMA de aluminio o acero inoxidable para controladores. Requieren certificado del fabricante.", "Alto"),
        ("Conductos Metálicos (RMC)", "Spec 624 / 830\n(Rigid Metallic Conduit)", "Hierro y Acero", "Tuberías conduit de acero galvanizado embebidas o expuestas. Sujetas a la norma 100% acero doméstico.", "Alto"),
        ("Tuberías Corrugadas de HDPE/PVC", "Spec 603 / 830\n(Plastic / HDPE Pipe)", "Material de Construcción", "Tuberías de polímeros sintéticos en sistemas de drenaje pluvial permanentes.", "Alto"),
        ("Rejillas, Marcos y Tapas Manhole", "Spec 604 / 908\n(Grates, Frames, Manholes)", "Hierro y Acero", "Fundiciones de hierro gris o hierro dúctil. Todo proceso de colado (casting) debe ser 100% en EE.UU.", "Alto"),
        ("Estructuras Prefabricadas (Vigas/Box)", "Spec 601 / 846 / 958 / 984\n(Precast Concrete Units)", "Producto Combinado", "Vigas pretensadas y alcantarillas de cajón. Requieren MTR del acero de refuerzo/torones y planta.", "Alto"),
        ("Juntas de Expansión", "Spec 606 / 837 / 939 / 942\n(Bridge Expansion Joints)", "Hierro y Acero / Combinado", "Juntas metálicas en puentes para movimientos térmicos. Trazabilidad completa del acero colado.", "Alto"),
        ("Bearings de Puentes", "Spec 606 / 979\n(Bridge Bearings)", "Producto Combinado", "Almohadillas de neopreno con placas internas de acero. Certificado de elastómero y MTR de placas.", "Alto"),
        ("Mallas Metálicas de Refuerzo", "Spec 602 / 922\n(Welded Wire Fabric)", "Hierro y Acero", "Malla electro-soldada en losas, aceras y pavimentos. Acero 100% fundido y estirado en EE.UU.", "Alto"),
        ("Cercas Permanentes (Chain Link)", "Spec 611 / 907 / 994\n(Chain Link Fence / Gate)", "Hierro y Acero", "Cercado de delimitación de servidumbre. Postes, alambres, tensores y galvanizado certificados.", "Alto"),
        ("Estructuras ITS (Pórticos/Mástiles)", "Spec 622 / 625 / 843\n(Gantries & ITS Towers)", "Hierro y Acero", "Pórticos de acero estructural para sensores y rótulos dinámicos. Fabricación 100% certificada.", "Alto"),
        ("Sensores de Tráfico / Loops", "Spec 625 / 833\n(Traffic Loop / Sensors)", "Producto Manufacturado", "Sensores embebidos en el pavimento o de radar. Declaración BABA de producto manufacturado.", "Medio"),
        ("Sistemas de Pesaje (WIM)", "Spec 625 / 981\n(Weigh In Motion WIM)", "Producto Manufacturado", "Placas piezoeléctricas o celdas de carga fijadas en carril. Certificación de componentes.", "Medio"),
        ("Equipos de Comunicación (Switches)", "Spec 625 / 831\n(Network Switch / Fiber)", "Producto Manufacturado", "Switches industriales de fibra óptica montados en gabinetes viales. Productos manufacturados de red.", "Medio"),
        ("Torres y Mástiles de Alumbrado", "Spec 624 / 970\n(High Mast Lighting Poles)", "Hierro y Acero", "Estructuras de acero reticuladas o monopolos. Cumplimiento total de acero nacional.", "Alto"),
        ("Componentes Prefabricados con Insertos", "Spec 601 / 990\n(MSE Wall Panels)", "Producto Combinado", "Pozos de registro y paneles MSE pre-colados. El acero e insertos metálicos requieren MTR.", "Alto"),
        ("Anclajes y Pernos de Anclaje", "Spec 605 / 949\n(Anchor Bolts ASTM F1554)", "Hierro y Acero", "Pernos de anclaje para postes y señales. Acero fundido, forjado y galvanizado en EE.UU.", "Alto"),
        ("Tachas Reflectivas Permanentes", "Spec 620 / 919\n(Raised Pavement Markers)", "Producto Manufacturado", "Marcadores de pavimento fijados con epóxico. Cuerpo y reflectores cumplen como producto manufacturado.", "Medio"),
        ("Pintura Termoplástica de Tráfico", "Spec 620 / 918\n(Thermoplastic Markings)", "Material de Construcción", "Demarcación permanente de pavimento. Resinas y microesferas de vidrio son materiales cubiertos.", "Medio"),
        ("Geotextiles y Mantas de Control", "Spec 628 / 828 / 924 / 929\n(Geotextile / Turf Mat)", "Material de Construcción", "Mantas sintéticas permanentes para estabilización de taludes. Productos de polímeros y plástico.", "Alto")
    ]

    for idx, row in enumerate(data_b):
        row_cells = table_b.add_row().cells
        bg_color = "F9FAFC" if idx % 2 == 1 else "FFFFFF"
        for i, val in enumerate(row):
            row_cells[i].width = col_widths_b[i]
            set_cell_background(row_cells[i], bg_color)
            set_cell_margins(row_cells[i], top=80, bottom=80, left=70, right=70)
            p = row_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(val)
            run.font.name = "Arial"
            run.font.size = Pt(8)
            if i == 1:
                run.bold = True
                run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
            elif i == 4:
                run.bold = True
                if val == "Alto":
                    run.font.color.rgb = RGBColor(0x9C, 0x00, 0x06)
                else:
                    run.font.color.rgb = RGBColor(0x9C, 0x65, 0x00)
            else:
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # --- SECCIÓN 4 ---
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    h4 = doc.add_heading(level=1)
    h4.paragraph_format.space_before = Pt(16)
    h4.paragraph_format.space_after = Pt(8)
    run_h4 = h4.add_run("4. SECCIÓN C – ÍTEMS QUE NORMALMENTE NO REQUIEREN CERTIFICADO DE MANUFACTURA")
    run_h4.bold = True
    run_h4.font.name = "Arial"
    run_h4.font.size = Pt(14)
    run_h4.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    p_secc = doc.add_paragraph()
    p_secc.paragraph_format.line_spacing = 1.15
    p_secc.paragraph_format.space_after = Pt(8)
    p_secc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_secc.add_run(
        "Las regulaciones de BABA (2 CFR 184.103) y la guía oficial de la FHWA excluyen expresamente de la obligación de "
        "presentar Certificados de Manufactura a aquellos equipos, herramientas, materiales y estructuras de carácter "
        "**temporal, consumible o de soporte operativo** (tales como la partida Spec 151 - Mobilization, Spec 983 - Office Trailers, etc.) "
        "que no pasen a formar parte física permanente de la infraestructura vial."
    )

    items_c = [
        ("Herramientas Manuales y Maquinaria del Contratista:", "Taladros, compresores, generadores portátiles, sierras de hormigón y herramientas menores del personal."),
        ("Equipo Pesado de Construcción (Spec 151 / Equipos):", "Excavadoras, grúas, pavimentadoras de asfalto/hormigón, rodillos compactadores, camiones de volteo. Medios de producción que se retiran del proyecto al finalizar."),
        ("Sistemas de Control de Tráfico Temporero (Spec 635 / 901 / 906):", "Conos de tráfico, tambores (drums), barreras de plástico llenas de agua y tableros de mensaje variable portátiles que se remueven tras la construcción."),
        ("Encofrados y Moldeados Temporeros (Spec 829 / Encofrados):", "Encofrados de madera o paneles de acero para colado de hormigón, puntales mecánicos y andamios que se remueven tras el curado."),
        ("Materiales Consumibles de Operación:", "Combustibles, lubricantes, aceites hidráulicos, trapos de limpieza, gas propano/acetileno para soldadura y EPP (cascos, chalecos, guantes)."),
        ("Estructuras y Oficinas de Campo Temporeras (Spec 983):", "Vagones u oficinas portátiles de inspección/contratista, cercas temporales y sanitarios portátiles.")
    ]

    for tit_c, desc_c in items_c:
        p_ic = doc.add_paragraph(style='List Bullet')
        p_ic.paragraph_format.space_after = Pt(4)
        run_bc = p_ic.add_run(tit_c + " ")
        run_bc.bold = True
        run_bc.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
        p_ic.add_run(desc_c)

    # --- SECCIÓN 5 ---
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    h5 = doc.add_heading(level=1)
    h5.paragraph_format.space_before = Pt(16)
    h5.paragraph_format.space_after = Pt(8)
    run_h5 = h5.add_run("5. CONCLUSIONES Y REGLA GENERAL DE AUDITORÍA BABA")
    run_h5.bold = True
    run_h5.font.name = "Arial"
    run_h5.font.size = Pt(14)
    run_h5.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    add_callout(
        doc,
        "REGLA GENERAL DE PRESUNCIÓN DE CUMPLIMIENTO BABA CON ESPECIFICACIONES ACT",
        "Como regla general de auditoría y gestión de contratos en proyectos de la FHWA y la ACT en Puerto Rico, "
        "TODO elemento que cumpla simultáneamente con las tres condiciones siguientes:\n\n"
        "1. Permanezca instalado de forma PERMANENTE en la obra terminada,\n"
        "2. Esté compuesto de HIERRO, ACERO, MATERIAL DE CONSTRUCCIÓN o PRODUCTO MANUFACTURADO (bajo partidas Spec 601, 602, 603, 604, 605, 606, 610, 611, 620, 622, 624, 625, 628, etc.), y\n"
        "3. Sea financiado en todo o en parte con FONDOS FEDERALES (FHWA / USDOT),\n\n"
        "DEBE PRESUMIRSE SUJETO A LA PRESENTACIÓN DE CERTIFICADO DE MANUFACTURA O CERTIFICACIÓN BABA, "
        "salvo que se cuente con una Exención Formal (Waiver) por escrito aprobada por la Administración Federal de Carreteras.",
        bg_hex="EBF1F5",
        border_hex="003366"
    )

    # --- SECCIÓN 6: ABOUT ---
    doc.add_page_break()
    h6 = doc.add_heading(level=1)
    h6.paragraph_format.space_before = Pt(16)
    h6.paragraph_format.space_after = Pt(8)
    run_h6 = h6.add_run("6. SECCIÓN ABOUT - CRÉDITOS DE AUTORÍA Y DESARROLLO")
    run_h6.bold = True
    run_h6.font.name = "Arial"
    run_h6.font.size = Pt(14)
    run_h6.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    add_callout(
        doc,
        "SOBRE EL AUTOR Y EL SISTEMA PACT",
        "Este documento técnico de análisis normativo ha sido diseñado y redactado por el Ing. Enrique Saavedra Sada, PE, "
        "creador y diseñador principal del programa PACT (Programa de Administración de Contratos).\n\n"
        "El sistema PACT integra de forma automatizada las validaciones de auditoría para Certificados de Manufactura, "
        "cruzando los códigos de especificación estándar de la ACT (Spec Codes) para garantizar que ninguna partida "
        "cubierta por BABA o Buy America pueda ser certificada para pago sin la evidencia documental correspondiente.\n\n"
        "Diseñador: Ing. Enrique Saavedra Sada, PE\n"
        "Especialidad: Administración de Contratos e Ingeniería de Construcción de Carreteras\n"
        "Fecha de Actualización: Julio 2026\n"
        "Documentación PACT - Puerto Rico",
        bg_hex="F4F6F9",
        border_hex="003366"
    )

    out_dir = r"C:\Users\Enrique Saavedra\Documents\PROGRAMAS AI\Programa ACT\Documentos"
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, "Analisis_Items_Certificado_BABA_FHWA.docx")
    doc.save(out_path)
    print(f"Documento BABA actualizado con éxito en: {out_path}")

if __name__ == "__main__":
    generate_baba_document()
