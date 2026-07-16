import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_manual():
    doc = Document()
    
    # Title
    title = doc.add_heading('Manual de Usuario Paso a Paso: PACT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Intro
    doc.add_paragraph('Bienvenido al manual paso a paso de PACT. Esta guía está diseñada para llevarte desde la creación de un proyecto nuevo hasta la gestión de documentos, certificaciones de pago y órdenes de cambio (CHO).')
    
    # Paso 1
    doc.add_heading('Paso 1: Comenzar un Proyecto Nuevo', level=1)
    doc.add_paragraph('1. Abre el programa PACT.\n'
                      '2. Dirígete a la sección de "Proyectos" en el menú principal.\n'
                      '3. Haz clic en el botón "Nuevo Proyecto".\n'
                      '4. Completa toda la información solicitada, como el nombre del proyecto, número de contrato, y cualquier otro dato básico.\n'
                      '5. Presiona "Guardar". ¡Listo! Tu proyecto ya está creado y listo para usarse.')
    
    # Paso 2
    doc.add_heading('Paso 2: Registrar los Documentos Necesarios', level=1)
    doc.add_paragraph('Es muy importante mantener la documentación al día.\n'
                      '1. Entra a tu proyecto recién creado y ve a la sección de "Documentos".\n'
                      '2. Haz clic en "Añadir" o "Registrar Documento".\n'
                      '3. Sube los documentos legales o requeridos para el proyecto.\n'
                      '4. Verifica cuidadosamente la fecha de expiración de cada documento. Registra únicamente los documentos que no han expirado aún. El sistema puede pedirte que ingreses la fecha de vigencia.\n'
                      '5. Guarda los registros.')

    # Paso 3
    doc.add_heading('Paso 3: Registrar Certificados de Manufactura', level=1)
    doc.add_paragraph('Antes de solicitar ciertos pagos, necesitarás comprobar los materiales.\n'
                      '1. Dirígete a la sección de "Certificados" o "Manufactura".\n'
                      '2. Selecciona "Registrar Certificado".\n'
                      '3. Por cada material utilizado, ingresa los datos correspondientes y adjunta el certificado del fabricante.\n'
                      '4. Asegúrate de registrar todos los certificados requeridos para evitar bloqueos en tus pagos futuros.')

    # Paso 4
    doc.add_heading('Paso 4: Hacer y Registrar 5 Certificaciones de Pago', level=1)
    doc.add_paragraph('A medida que el proyecto avanza, deberás cobrar por el trabajo realizado.\n'
                      '1. Ve a la pestaña de "Certificaciones de Pago".\n'
                      '2. Haz clic en "Nueva Certificación".\n'
                      '3. Certificación #1: Ingresa el porcentaje de avance o la cantidad trabajada en este primer periodo. Verifica que los certificados de manufactura requeridos estén en orden y guarda la certificación.\n'
                      '4. Certificaciones #2, #3, #4 y #5: Conforme avance el tiempo, repetirás el proceso (pasos 2 y 3) para cada una de las siguientes 4 certificaciones. El programa PACT restará automáticamente lo que ya has cobrado en certificaciones anteriores para mostrarte tu balance correcto.')

    # Paso 5
    doc.add_heading('Paso 5: Hacer Dos Órdenes de Cambio (CHO)', level=1)
    doc.add_paragraph('A veces los proyectos necesitan modificaciones (más tiempo o dinero).\n'
                      '1. Navega a la sección de "Órdenes de Cambio" (CHO).\n'
                      '2. Haz clic en "Nuevo CHO".\n'
                      '3. Primer CHO: Detalla qué cambió (por ejemplo, trabajo adicional), pon los nuevos montos y justifica el cambio. Guarda el documento.\n'
                      '4. Segundo CHO: Repite el proceso para registrar una segunda orden de cambio.\n'
                      '5. Una vez guardados y aprobados, estos CHO modificarán el monto total de tu contrato y se reflejarán en tus próximas certificaciones de pago.')

    # Nueva seccion: Resumen
    doc.add_heading('Sección de Resumen (Dashboard)', level=1)
    doc.add_paragraph('La sección de "Resumen" te permite tener una visión general del estado actual de tu proyecto. Aquí podrás observar:\n'
                      '• El monto total del contrato original y el monto actual ajustado con los CHO aprobados.\n'
                      '• El total facturado hasta el momento a través de tus certificaciones de pago.\n'
                      '• El balance restante disponible en el contrato.\n'
                      '• El estado de tus documentos (cuáles están vigentes y cuáles están próximos a expirar).\n'
                      '• Alertas sobre certificados de manufactura pendientes.\n'
                      'Es una herramienta excelente para monitorear la salud financiera y administrativa de tu proyecto en un solo vistazo.')
    
    # Save document
    # Ensure no .docx extension is repeated if provided, but the user asked for this exact path, adding .docx
    file_path = r"C:\Users\Enrique Saavedra\Documents\PROGRAMAS AI\Programa ACT\Documentos\MANUAL PACT JULIO 2026.docx"
    doc.save(file_path)
    print(f"Document saved to {file_path}")

if __name__ == "__main__":
    create_manual()
